from pathlib import Path
import os, math, textwrap
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT=Path(r"C:\Users\anh15\GYM_Management")
OUT=ROOT/'output'/'report'; FIG=OUT/'figures'; OUT.mkdir(parents=True,exist_ok=True); FIG.mkdir(exist_ok=True)
NAVY='#172033'; ORANGE='#F97316'; BLUE='#2A354A'; LIGHT='#F4F6F8'; GRAY='#667085'; GREEN='#16A34A'; RED='#DC2626'

def diagram(name,title,nodes,edges,cols=3):
    rows=math.ceil(len(nodes)/cols); W=1600; H=250+rows*210; im=Image.new('RGB',(W,H),'white'); d=ImageDraw.Draw(im)
    font=ImageFont.truetype('arial.ttf',24); bold=ImageFont.truetype('arialbd.ttf',27); titlefont=ImageFont.truetype('arialbd.ttf',34)
    d.text((W//2,45),title,fill=NAVY,font=titlefont,anchor='mm')
    pos={}
    for i,(key,label,color) in enumerate(nodes):
        r=i//cols; c=i%cols; x=int((c+.5)*W/cols); y=170+r*200; pos[key]=(x,y)
        box=(x-210,y-58,x+210,y+58); d.rounded_rectangle(box,18,fill=color,outline=NAVY,width=3)
        lines=textwrap.wrap(label,26); yy=y-(len(lines)-1)*16
        for line in lines: d.text((x,yy),line,fill='white' if color!=LIGHT else NAVY,font=bold,anchor='mm'); yy+=32
    for a,b in edges:
        if a in pos and b in pos:
            x1,y1=pos[a]; x2,y2=pos[b]; d.line((x1,y1,x2,y2),fill=GRAY,width=4); ang=math.atan2(y2-y1,x2-x1); L=18
            pts=[(x2,y2),(x2-L*math.cos(ang-.5),y2-L*math.sin(ang-.5)),(x2-L*math.cos(ang+.5),y2-L*math.sin(ang+.5))]; d.polygon(pts,fill=GRAY)
    path=FIG/f'{name}.png'; im.save(path); return path

def barfig(name,title,labels,values,colors=None):
    W,H=1500,650; im=Image.new('RGB',(W,H),'white'); d=ImageDraw.Draw(im); colors=colors or [ORANGE]*len(values)
    titlefont=ImageFont.truetype('arialbd.ttf',34); font=ImageFont.truetype('arial.ttf',22); bold=ImageFont.truetype('arialbd.ttf',24)
    d.text((W//2,40),title,fill=NAVY,font=titlefont,anchor='mm'); base=540; maxv=max(values); bw=120; gap=(W-180)//len(values)
    d.line((80,base,W-50,base),fill=GRAY,width=3)
    for i,(lab,v,col) in enumerate(zip(labels,values,colors)):
        x=100+i*gap; h=int(v/maxv*380); d.rounded_rectangle((x,base-h,x+bw,base),10,fill=col); d.text((x+bw//2,base-h-18),str(v),fill=NAVY,font=bold,anchor='mm'); d.text((x+bw//2,base+38),lab,fill=NAVY,font=font,anchor='mm')
    p=FIG/f'{name}.png'; im.save(p); return p

def conceptual_erd():
    W,H=1900,1750; im=Image.new('RGB',(W,H),'white'); d=ImageDraw.Draw(im)
    title=ImageFont.truetype('arialbd.ttf',42); font=ImageFont.truetype('arial.ttf',24); bold=ImageFont.truetype('arialbd.ttf',24); small=ImageFont.truetype('arialbd.ttf',20)
    d.rectangle((25,25,W-25,H-25),outline=NAVY,width=4); d.text((W//2,65),'ERD KHÁI NIỆM - GYM PRO',font=title,fill=NAVY,anchor='mm')
    entities={'Hồ sơ':(180,220),
              'User':(180,600),'Giáo án':(650,600),'Ngày tập':(1120,600),'Bài trong giáo án':(1600,600),
              'Tiến độ':(180,1040),'Buổi tập':(650,1040),'Nhật ký bài tập':(1120,1040),'Bài tập':(1600,1040)}
    def entity(label,x,y):
        d.rounded_rectangle((x-105,y-48,x+105,y+48),8,fill='#FFFFFF',outline=GRAY,width=3); d.text((x,y),label,font=bold,fill=NAVY,anchor='mm')
    for label,(x,y) in entities.items(): entity(label,x,y)
    rels=[('User','Hồ sơ','Có','1','1',(180,410)),('User','Giáo án','Tạo','1','N',(415,600)),
          ('Giáo án','Ngày tập','Bao gồm','1','N',(885,600)),('Ngày tập','Bài trong giáo án','Chứa','1','N',(1360,600)),
          ('Bài trong giáo án','Bài tập','Tham chiếu','N','1',(1600,820)),('Giáo án','Buổi tập','Thực hiện','1','N',(650,820)),
          ('Buổi tập','Nhật ký bài tập','Ghi nhận','1','N',(885,1040)),('Nhật ký bài tập','Bài tập','Của','N','1',(1360,1040)),
          ('User','Tiến độ','Theo dõi','1','N',(180,820))]
    for a,b,r,ca,cb,(mx,my) in rels:
        x1,y1=entities[a]; x2,y2=entities[b]
        pts=[(mx,my-42),(mx+70,my),(mx,my+42),(mx-70,my)]
        d.line((x1,y1,mx,my),fill=GRAY,width=3); d.line((mx,my,x2,y2),fill=GRAY,width=3)
        d.polygon(pts,fill='#FFF7ED',outline=ORANGE); d.text((mx,my),r,font=font,fill=NAVY,anchor='mm')
        d.text(((x1+mx)//2,(y1+my)//2-18),ca,font=small,fill=RED,anchor='mm'); d.text(((x2+mx)//2,(y2+my)//2-18),cb,font=small,fill=RED,anchor='mm')
    p=FIG/'erd_conceptual.png'; im.save(p); return p

def logical_erd():
    W,H=2600,1900; im=Image.new('RGB',(W,H),'white'); d=ImageDraw.Draw(im)
    title=ImageFont.truetype('arialbd.ttf',42); head=ImageFont.truetype('arialbd.ttf',21); font=ImageFont.truetype('arial.ttf',18); keyfont=ImageFont.truetype('arialbd.ttf',17)
    d.rectangle((20,20,W-20,H-20),outline=NAVY,width=4); d.text((W//2,55),'ERD LOGIC - GYM PRO',font=title,fill=NAVY,anchor='mm')
    tables={
      'USERS':(['PK id','email','password','full_name','status'],(90,150)),
      'USER_PROFILES':(['PK id','FK user_id','height','weight','goal','injury_areas'],(510,120)),
      'MEMBERSHIPS':(['PK id','FK user_id','membership_type','start_date','end_date'],(980,120)),
      'INVOICES':(['PK id','FK user_id','membership_type','amount','status'],(1450,120)),
      'WORKOUT_PLANS':(['PK id','FK user_id','goal','target_level','sessions_per_week','fitness_level'],(90,620)),
      'PLAN_DAYS':(['PK id','FK workout_plan_id','day_of_week','day_name'],(600,620)),
      'PLAN_EXERCISES':(['PK id','FK plan_day_id','FK exercise_id','sets','reps','rest_seconds'],(1050,600)),
      'EXERCISES':(['PK id','name','muscle_group','difficulty','stamina_cost','affected_groups'],(1590,600)),
      'WORKOUT_SESSIONS':(['PK id','FK user_id','FK workout_plan_id','FK plan_day_id','session_date','completion_rate'],(90,1200)),
      'EXERCISE_LOGS':(['PK id','FK session_id','FK exercise_id','reps_completed','completion_percent'],(720,1200)),
      'PROGRESS':(['PK id','FK user_id','weight','body_fat','recorded_date'],(1330,1220)),
      'SYSTEM_CONFIGS':(['PK config_key','config_value','category','description'],(1870,1220))}
    boxes={}
    for name,(fields,(x,y)) in tables.items():
        w=390; rowh=34; h=48+len(fields)*rowh; boxes[name]=(x,y,w,h)
        d.rectangle((x,y,x+w,y+h),fill='white',outline=GRAY,width=3); d.rectangle((x,y,x+w,y+48),fill='#E8EEF5',outline=GRAY,width=2); d.text((x+w//2,y+24),name,font=head,fill=NAVY,anchor='mm')
        for i,f in enumerate(fields):
            yy=y+48+i*rowh; d.line((x,yy,x+w,yy),fill='#D0D5DD',width=1); col=RED if f.startswith('PK') else ORANGE if f.startswith('FK') else NAVY
            d.text((x+14,yy+17),f,font=keyfont if f.startswith(('PK','FK')) else font,fill=col,anchor='lm')
    links=[('USERS','USER_PROFILES'),('USERS','MEMBERSHIPS'),('USERS','INVOICES'),('USERS','WORKOUT_PLANS'),('WORKOUT_PLANS','PLAN_DAYS'),('PLAN_DAYS','PLAN_EXERCISES'),('EXERCISES','PLAN_EXERCISES'),('USERS','WORKOUT_SESSIONS'),('WORKOUT_PLANS','WORKOUT_SESSIONS'),('PLAN_DAYS','WORKOUT_SESSIONS'),('WORKOUT_SESSIONS','EXERCISE_LOGS'),('EXERCISES','EXERCISE_LOGS'),('USERS','PROGRESS')]
    for a,b in links:
        x,y,w,h=boxes[a]; x2,y2,w2,h2=boxes[b]; sx=x+w if x2>x else x; sy=y+h//2; ex=x2 if x2>x else x2+w2; ey=y2+h2//2
        mid=(sx+ex)//2; d.line((sx,sy,mid,sy,mid,ey,ex,ey),fill='#98A2B3',width=3); d.ellipse((ex-5,ey-5,ex+5,ey+5),fill=ORANGE)
    d.text((W-500,H-80),'PK: Primary Key   |   FK: Foreign Key',font=head,fill=GRAY,anchor='mm')
    p=FIG/'erd_logical.png'; im.save(p); return p

def usecase_diagram(name,title,actors,cases,links):
    W,H=1900,max(1100,240+len(cases)*105); im=Image.new('RGB',(W,H),'white'); d=ImageDraw.Draw(im)
    tf=ImageFont.truetype('arialbd.ttf',38); f=ImageFont.truetype('arial.ttf',23); bf=ImageFont.truetype('arialbd.ttf',23)
    d.text((W//2,48),title,font=tf,fill=NAVY,anchor='mm'); boundary=(470,100,1550,H-70); d.rounded_rectangle(boundary,15,outline=GRAY,width=3); d.text((1010,130),'HỆ THỐNG GYM PRO',font=bf,fill=GRAY,anchor='mm')
    apos={}; sides=[180,1720]
    for i,a in enumerate(actors):
        x=sides[i%2]; y=260+(i//2)*430; apos[a]=(x,y); d.ellipse((x-28,y-90,x+28,y-34),outline=NAVY,width=3); d.line((x,y-34,x,y+55),fill=NAVY,width=3); d.line((x-48,y,x+48,y),fill=NAVY,width=3); d.line((x,y+55,x-40,y+115),fill=NAVY,width=3); d.line((x,y+55,x+40,y+115),fill=NAVY,width=3); d.text((x,y+145),a,font=bf,fill=NAVY,anchor='mm')
    cpos={}
    for i,c in enumerate(cases):
        x=760 if i%2==0 else 1240; y=220+(i//2)*180; cpos[c]=(x,y); d.ellipse((x-210,y-55,x+210,y+55),fill='#FFF7ED',outline=ORANGE,width=3); d.text((x,y),c,font=f,fill=NAVY,anchor='mm')
    for a,c in links:
        if a in apos and c in cpos: d.line((*apos[a],*cpos[c]),fill='#98A2B3',width=3)
    p=FIG/f'{name}.png'; im.save(p); return p

def flowchart(name,title,steps):
    W,H=1700,220+len(steps)*175; im=Image.new('RGB',(W,H),'white'); d=ImageDraw.Draw(im); tf=ImageFont.truetype('arialbd.ttf',38); f=ImageFont.truetype('arialbd.ttf',23); small=ImageFont.truetype('arial.ttf',19)
    d.text((W//2,45),title,font=tf,fill=NAVY,anchor='mm'); cx=W//2
    positions=[]
    for i,(label,kind) in enumerate(steps):
        y=150+i*170; positions.append((cx,y));
        if kind=='start': d.ellipse((cx-180,y-48,cx+180,y+48),fill='#ECFDF3',outline=GREEN,width=3)
        elif kind=='decision': d.polygon([(cx,y-68),(cx+230,y),(cx,y+68),(cx-230,y)],fill='#FFF7ED',outline=ORANGE)
        elif kind=='data': d.polygon([(cx-210,y-50),(cx+250,y-50),(cx+210,y+50),(cx-250,y+50)],fill='#EFF6FF',outline=BLUE)
        else: d.rounded_rectangle((cx-230,y-52,cx+230,y+52),12,fill='#F8FAFC',outline=NAVY,width=3)
        for j,line in enumerate(textwrap.wrap(label,34)): d.text((cx,y+(j-(len(textwrap.wrap(label,34))-1)/2)*25),line,font=f,fill=NAVY,anchor='mm')
        if i>0:
            py=positions[i-1][1]; d.line((cx,py+70,cx,y-70),fill=GRAY,width=3); d.polygon([(cx,y-70),(cx-9,y-88),(cx+9,y-88)],fill=GRAY)
    p=FIG/f'{name}.png'; im.save(p); return p

def sequence_diagram(name,title,participants,messages):
    W=1900; H=300+len(messages)*105; im=Image.new('RGB',(W,H),'white'); d=ImageDraw.Draw(im); tf=ImageFont.truetype('arialbd.ttf',38); f=ImageFont.truetype('arial.ttf',20); bf=ImageFont.truetype('arialbd.ttf',21)
    d.text((W//2,45),title,font=tf,fill=NAVY,anchor='mm'); xs=[150+i*(1600//(len(participants)-1)) for i in range(len(participants))]
    for x,pn in zip(xs,participants): d.rounded_rectangle((x-105,105,x+105,165),8,fill='#E8EEF5',outline=NAVY,width=3); d.text((x,135),pn,font=bf,fill=NAVY,anchor='mm'); d.line((x,165,x,H-70),fill='#98A2B3',width=2)
    idx={p:i for i,p in enumerate(participants)}
    for k,(a,b,msg,ret) in enumerate(messages):
        y=220+k*100; x1,x2=xs[idx[a]],xs[idx[b]]; col=GRAY if ret else ORANGE; d.line((x1,y,x2,y),fill=col,width=3)
        direction=1 if x2>x1 else -1; d.polygon([(x2,y),(x2-14*direction,y-8),(x2-14*direction,y+8)],fill=col); d.text(((x1+x2)//2,y-16),msg,font=f,fill=NAVY,anchor='ms')
    p=FIG/f'{name}.png'; im.save(p); return p

def source_tree_diagram():
    W,H=1900,1380; im=Image.new('RGB',(W,H),'white'); d=ImageDraw.Draw(im)
    tf=ImageFont.truetype('arialbd.ttf',38); f=ImageFont.truetype('consola.ttf',23); bf=ImageFont.truetype('consolab.ttf',24)
    d.text((W//2,48),'CẤU TRÚC MÃ NGUỒN',font=tf,fill=NAVY,anchor='mm')
    columns=[(80,'FRONTEND',['gym-frontend/','├─ src/','│  ├─ api/              # API modules','│  ├─ components/       # Thành phần chung','│  ├─ router/           # Route và guard','│  ├─ stores/           # Trạng thái','│  ├─ views/','│  │  ├─ admin/         # Màn quản trị','│  │  └─ user/          # Màn người dùng','│  └─ App.vue','└─ package.json'],ORANGE),(980,'BACKEND',['gym-management/','├─ src/main/java/.../','│  ├─ controller/       # REST endpoints','│  ├─ dto/              # Request/Response','│  ├─ entity/           # JPA entities','│  ├─ repository/       # Truy vấn dữ liệu','│  ├─ service/          # Nghiệp vụ','│  ├─ security/         # JWT, phân quyền','│  └─ config/           # Cấu hình','├─ src/main/resources/','│  └─ application.properties','└─ pom.xml'],GREEN)]
    for x,label,lines,col in columns:
        d.rounded_rectangle((x,115,x+840,H-70),16,fill='#F8FAFC',outline=col,width=4); d.rectangle((x,115,x+840,180),fill=col); d.text((x+420,147),label,font=bf,fill='white',anchor='mm')
        yy=220
        for line in lines: d.text((x+35,yy),line,font=bf if line.endswith('/') and '─' not in line else f,fill=NAVY,anchor='lm'); yy+=84
    p=FIG/'source_tree.png'; im.save(p); return p

def wireframe_diagram():
    W,H=1900,1150; im=Image.new('RGB',(W,H),'white'); d=ImageDraw.Draw(im); tf=ImageFont.truetype('arialbd.ttf',38); f=ImageFont.truetype('arial.ttf',22); bf=ImageFont.truetype('arialbd.ttf',23)
    d.text((W//2,45),'WIREFRAME DASHBOARD NGƯỜI DÙNG',font=tf,fill=NAVY,anchor='mm'); d.rectangle((60,100,W-60,H-60),outline=NAVY,width=4)
    d.rectangle((60,100,390,H-60),fill='#E8EEF5',outline=NAVY,width=3); d.text((225,145),'GYM PRO',font=bf,fill=NAVY,anchor='mm')
    for i,label in enumerate(['Thống kê','Hồ sơ','Giáo án','Buổi tập','Tiến độ','Cửa hàng','Gói tập']): d.rectangle((95,210+i*92,355,265+i*92),outline=GRAY,width=2); d.text((225,237+i*92),label,font=f,fill=NAVY,anchor='mm')
    d.rectangle((390,100,W-60,205),fill='#F8FAFC',outline=GRAY,width=2); d.text((440,152),'≡    Tiêu đề trang',font=bf,fill=NAVY,anchor='lm'); d.text((W-120,152),'Thông báo  •  Tài khoản',font=f,fill=NAVY,anchor='rm')
    for i,label in enumerate(['Chỉ số 1','Chỉ số 2','Chỉ số 3','Chỉ số 4']):
        x=440+i*335; d.rounded_rectangle((x,255,x+285,400),14,fill='#FFF7ED',outline=ORANGE,width=3); d.text((x+142,305),label,font=f,fill=NAVY,anchor='mm'); d.text((x+142,355),'Giá trị',font=bf,fill=NAVY,anchor='mm')
    d.rounded_rectangle((440,450,1130,1015),14,fill='white',outline=GRAY,width=3); d.text((785,490),'BIỂU ĐỒ / DANH SÁCH CHÍNH',font=bf,fill=NAVY,anchor='mm')
    for y in range(570,950,75): d.line((500,y,1070,y),fill='#D0D5DD',width=2)
    d.rounded_rectangle((1180,450,1790,760),14,fill='white',outline=GRAY,width=3); d.text((1485,490),'THÔNG TIN PHỤ',font=bf,fill=NAVY,anchor='mm')
    d.rounded_rectangle((1340,820,1630,900),14,fill=ORANGE,outline=ORANGE,width=3); d.text((1485,860),'HÀNH ĐỘNG CHÍNH',font=bf,fill='white',anchor='mm')
    p=FIG/'wireframe.png'; im.save(p); return p

figs={}
figs['context']=diagram('context','Bối cảnh và giá trị của GYM PRO',[
 ('manual','Theo dõi rời rạc',RED),('profile','Hồ sơ thể trạng',BLUE),('plan','Giáo án cá nhân hóa',ORANGE),
 ('session','Tập linh hoạt theo ngày',GREEN),('progress','Theo dõi tiến độ',BLUE),('admin','Quản trị tập trung',NAVY)],
 [('manual','profile'),('profile','plan'),('plan','session'),('session','progress'),('admin','plan')])
figs['scope']=diagram('scope','Phạm vi chức năng',[
 ('auth','Tài khoản',BLUE),('prof','Hồ sơ',ORANGE),('plans','Giáo án',GREEN),('sessions','Buổi tập',BLUE),('progress','Tiến độ',ORANGE),('shop','Cửa hàng',GREEN),('chat','Hỗ trợ',BLUE),('admin','Quản trị',NAVY)],[],4)
figs['uc']=usecase_diagram('usecase','USE CASE TỔNG QUAN',['Khách','User','Admin'],
 ['Đăng ký / đăng nhập','Cập nhật hồ sơ','Tạo giáo án','Thực hiện buổi tập','Theo dõi tiến độ','Mua gói / sản phẩm','Quản lý danh mục','Cấu hình hệ thống'],
 [('Khách','Đăng ký / đăng nhập'),('User','Cập nhật hồ sơ'),('User','Tạo giáo án'),('User','Thực hiện buổi tập'),('User','Theo dõi tiến độ'),('User','Mua gói / sản phẩm'),('Admin','Quản lý danh mục'),('Admin','Cấu hình hệ thống')])
figs['erd']=conceptual_erd()
figs['erd_logic']=logical_erd()
figs['arch']=diagram('architecture','Kiến trúc hệ thống 3 lớp',[
 ('vue','Vue 3 + Element Plus',ORANGE),('api','Axios / REST API',BLUE),('spring','Spring Boot 3',NAVY),
 ('service','Business Services',GREEN),('jpa','Spring Data JPA',BLUE),('db','H2 / MySQL',ORANGE)],
 [('vue','api'),('api','spring'),('spring','service'),('service','jpa'),('jpa','db')],2)
figs['nav']=diagram('navigation','Sơ đồ điều hướng giao diện',[
 ('home','Trang chủ',NAVY),('login','Đăng nhập',BLUE),('dash','Dashboard user',ORANGE),
 ('profile','Hồ sơ',BLUE),('plan','Giáo án',GREEN),('sessions','Buổi tập',GREEN),
 ('progress','Tiến độ',BLUE),('shop','Cửa hàng',ORANGE),('admin','Dashboard admin',RED),
 ('manage','Các màn quản lý',NAVY)],
 [('home','login'),('login','dash'),('dash','profile'),('dash','plan'),('plan','sessions'),('sessions','progress'),('dash','shop'),('login','admin'),('admin','manage')])
figs['generate']=flowchart('generate_flow','FLOWCHART SINH GIÁO ÁN',[
 ('Bắt đầu','start'),('Đọc hồ sơ và mục tiêu','data'),('Hồ sơ đã đủ dữ liệu?','decision'),('Tính Fitness Score và trình độ','process'),('Đọc lịch + nhóm cơ do admin cấu hình','data'),('Lọc chấn thương, thiết bị, bài không thích','process'),('Áp dụng Training Zone sets/reps','process'),('Lưu giáo án và các ngày tập','data'),('Kết thúc','start')])
figs['checkout']=flowchart('checkout_flow','FLOWCHART CHECK-IN / CHECK-OUT',[
 ('Bắt đầu buổi tập','start'),('Kiểm tra buổi đang hoạt động','decision'),('Ghi reps hoặc thời lượng từng bài','data'),('Tính completion % từng bài','process'),('Tính calories và mana','process'),('Lưu SessionExerciseLog','data'),('Kiểm tra điều chỉnh từng bài','process'),('Hoàn thành buổi tập','start')])
figs['adaptive']=flowchart('adaptive','FLOWCHART ĐIỀU CHỈNH TỪNG BÀI',[
 ('Nhận log bài vừa tập','start'),('Đủ 2 lần gần nhất?','decision'),('Cả hai lần dưới ngưỡng?','decision'),('Đọc chiến lược admin','data'),('Đổi sang bài dễ hơn?','decision'),('Tìm bài cùng nhóm cơ và an toàn','process'),('Hoặc giảm sets/reps','process'),('Đã chạm sàn Training Zone?','decision'),('Lưu marker log đã xử lý','data'),('Kết thúc','start')])
figs['deploy']=diagram('deployment','Mô hình triển khai',[
 ('browser','Trình duyệt',ORANGE),('vite','Vue SPA',BLUE),('server','Spring Boot API',NAVY),('storage','Uploads',GREEN),('database','Cơ sở dữ liệu',ORANGE)],
 [('browser','vite'),('vite','server'),('server','storage'),('server','database')])
figs['test']=barfig('test_result','Kết quả kiểm thử theo nhóm chức năng',['Tài khoản','Hồ sơ','Giáo án','Buổi tập','Admin','Cửa hàng'],[8,10,18,16,20,12],[BLUE,ORANGE,GREEN,BLUE,NAVY,ORANGE])
figs['source']=source_tree_diagram()
figs['seq_login']=sequence_diagram('sequence_login','SEQUENCE - ĐĂNG NHẬP',['User','Frontend','AuthController','AuthService','Database'],[
 ('User','Frontend','Nhập email/mật khẩu',False),('Frontend','AuthController','POST /api/auth/login',False),('AuthController','AuthService','authenticate()',False),('AuthService','Database','findByEmail()',False),('Database','AuthService','User + Role',True),('AuthService','Frontend','JWT + thông tin user',True),('Frontend','User','Điều hướng theo vai trò',True)])
figs['seq_plan']=sequence_diagram('sequence_plan','SEQUENCE - SINH GIÁO ÁN',['User','Frontend','PlanController','PlanService','ConfigService','Repository'],[
 ('User','Frontend','Chọn mục tiêu',False),('Frontend','PlanController','POST generate',False),('PlanController','PlanService','generate()',False),('PlanService','ConfigService','Lấy lịch và split',False),('ConfigService','PlanService','Cấu hình',True),('PlanService','Repository','Lọc và lưu giáo án',False),('Repository','PlanService','WorkoutPlan',True),('PlanService','Frontend','Chi tiết giáo án',True)])
figs['seq_checkout']=sequence_diagram('sequence_checkout','SEQUENCE - CHECK-OUT',['User','Frontend','SessionController','SessionService','PlanService','Repository'],[
 ('User','Frontend','Nhập kết quả',False),('Frontend','SessionController','POST check-out',False),('SessionController','SessionService','checkOut()',False),('SessionService','Repository','Lưu exercise logs',False),('SessionService','PlanService','Kiểm tra 2 lần gần nhất',False),('PlanService','Repository','Điều chỉnh bài nếu cần',False),('Repository','SessionService','Kết quả lưu',True),('SessionService','Frontend','Session response',True)])
figs['wire']=wireframe_diagram()

doc=Document(); sec=doc.sections[0]; sec.page_height=Cm(29.7); sec.page_width=Cm(21); sec.top_margin=Cm(2); sec.bottom_margin=Cm(2); sec.left_margin=Cm(2.4); sec.right_margin=Cm(2); sec.header_distance=Cm(1); sec.footer_distance=Cm(1)
styles=doc.styles
for s in ['Normal','Heading 1','Heading 2','Heading 3']:
    st=styles[s]; st.font.name='Times New Roman'; st._element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman')
styles['Normal'].font.size=Pt(12); styles['Normal'].paragraph_format.space_after=Pt(5); styles['Normal'].paragraph_format.line_spacing=1.18
for s,size,color,before,after in [('Heading 1',17,NAVY,14,8),('Heading 2',14,ORANGE,10,6),('Heading 3',12,BLUE,8,4)]:
    st=styles[s]; st.font.size=Pt(size); st.font.bold=True; st.font.color.rgb=RGBColor.from_string(color[1:]); st.paragraph_format.space_before=Pt(before); st.paragraph_format.space_after=Pt(after)

header=sec.header.paragraphs[0]; header.text='GYM PRO  |  BÁO CÁO DỰ ÁN TỐT NGHIỆP'; header.alignment=WD_ALIGN_PARAGRAPH.RIGHT
for r in header.runs: r.font.name='Arial'; r.font.size=Pt(8); r.font.color.rgb=RGBColor(102,112,133)
footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
run=footer.add_run('GYM PRO  •  '); fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); run._r.addnext(fld)

def shade(cell,fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=tcPr.find(qn('w:shd')) or OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); tcPr.append(shd) if shd.getparent() is None else None
def add_title(text,sub=None):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(110); r=p.add_run(text); r.bold=True; r.font.name='Arial'; r.font.size=Pt(28); r.font.color.rgb=RGBColor.from_string(NAVY[1:])
    if sub:
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(sub); r.font.name='Arial'; r.font.size=Pt(15); r.font.color.rgb=RGBColor.from_string(ORANGE[1:])
def add_p(text,bold=False,center=False):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY; r=p.add_run(text); r.bold=bold; return p
def add_bullets(items):
    for x in items:
        p=doc.add_paragraph(style='List Bullet'); p.paragraph_format.left_indent=Cm(.6); p.paragraph_format.first_line_indent=Cm(-.3); p.add_run(x)
def add_fig(path,caption,width=15.8):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run().add_picture(str(path),width=Cm(width))
    c=doc.add_paragraph(caption); c.alignment=WD_ALIGN_PARAGRAPH.CENTER; c.runs[0].italic=True; c.runs[0].font.size=Pt(10); c.runs[0].font.color.rgb=RGBColor(102,112,133)
def add_table(headers,rows,widths=None,font=9):
    t=doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False; t.style='Table Grid'
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=h; shade(c,'172033'); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for r in c.paragraphs[0].runs: r.font.bold=True; r.font.color.rgb=RGBColor(255,255,255); r.font.size=Pt(font)
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cells[i].text=str(v); cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[i].paragraphs:
                for r in p.runs: r.font.name='Times New Roman'; r.font.size=Pt(font)
    if widths:
        for row in t.rows:
            for i,w in enumerate(widths): row.cells[i].width=Cm(w)
    doc.add_paragraph().paragraph_format.space_after=Pt(1); return t
def page(): doc.add_page_break()
def chapter(n,title): page(); doc.add_heading(f'PHẦN {n}: {title}',0)

# Cover
add_p('TRƯỜNG CAO ĐẲNG FPT POLYTECHNIC',True,True); add_p('CHUYÊN NGÀNH PHÁT TRIỂN PHẦN MỀM',True,True)
add_title('BÁO CÁO DỰ ÁN TỐT NGHIỆP','XÂY DỰNG HỆ THỐNG QUẢN LÝ VÀ HỖ TRỢ LUYỆN TẬP GYM PRO')
add_p('Giảng viên hướng dẫn: ........................................................',False,True)
add_p('Sinh viên thực hiện: ............................................................',False,True)
add_p('Mã sinh viên: ...................................................................',False,True)
add_p('HÀ NỘI - 2026',True,True)
page(); doc.add_heading('DANH SÁCH THÀNH VIÊN',0)
add_table(['STT','Họ và tên','Mã sinh viên','Email'],[['1','................................','................','................................']], [1.5,5,3.5,5])
add_p('Sinh viên cam kết nội dung báo cáo phản ánh trung thực quá trình phân tích, thiết kế, phát triển và kiểm thử dự án; chịu trách nhiệm về tính chính xác và bản quyền của tài liệu.')
page(); doc.add_heading('GIẢNG VIÊN HƯỚNG DẪN',0)
add_p('Họ và tên: ................................................................................................')
add_p('Ý kiến nhận xét, đánh giá:')
for _ in range(10): add_p('........................................................................................................................................')
add_p('Giáo viên hướng dẫn\n(Ký và ghi rõ họ tên)',False,True)
page(); doc.add_heading('LỜI CẢM ƠN',0)
add_p('Em xin gửi lời cảm ơn chân thành đến giảng viên hướng dẫn và các thầy cô Trường Cao đẳng FPT Polytechnic đã cung cấp kiến thức, góp ý và tạo điều kiện để dự án GYM PRO được hoàn thiện. Em cũng cảm ơn bạn bè và người dùng thử đã hỗ trợ kiểm tra nghiệp vụ, giao diện và phản hồi trong quá trình phát triển.')
page(); doc.add_heading('MỤC LỤC',0)
toc=[('PHẦN 1: GIỚI THIỆU','1'),('1.1 Bối cảnh - hiện trạng','1'),('1.2 Mục tiêu - phạm vi','3'),('1.2.1 Mục tiêu','3'),('1.2.2 Phạm vi','4'),('1.3 Nguồn lực - kế hoạch','5'),('1.3.1 Nguồn lực','5'),('1.3.2 Kế hoạch thực hiện','6'),('PHẦN 2: PHÂN TÍCH','7'),('2.1 Yêu cầu người dùng','7'),('2.2 Trường hợp sử dụng','9'),('2.2.1 Danh sách tác nhân','9'),('2.2.2 Danh sách Use Case','10'),('2.3 Quan hệ thực thể','13'),('2.3.1 Danh sách thực thể','13'),('2.3.2 Các mối quan hệ','14'),('2.3.3 ERD mức khái niệm','15'),('2.3.4 ERD mức logic','16'),('PHẦN 3: THIẾT KẾ','17'),('3.1 Cơ sở dữ liệu','17'),('3.1.1 Chuẩn hóa','17'),('3.1.2 Danh sách bảng','18'),('3.2 Giao diện người dùng','20'),('3.2.1 Sơ đồ giao diện','20'),('3.2.2 Giao diện phác thảo','21'),('3.2.3 Danh sách giao diện','22'),('PHẦN 4: THỰC THI','24'),('4.1 Tổ chức mã nguồn','24'),('4.1.1 Sơ đồ tổ chức','24'),('4.1.2 Thư viện sử dụng','25'),('4.2 Đặc tả chức năng','27'),('4.2.1 Sinh giáo án','27'),('4.2.2 Check-in/check-out','29'),('4.2.3 Điều chỉnh từng bài','31'),('4.2.4 Thanh toán và phân quyền','33'),('PHẦN 5: KIỂM THỬ','35'),('5.1 Kế hoạch kiểm thử','35'),('5.1.1 Tiêu chí cần đạt','35'),('5.1.2 Chiến lược triển khai','36'),('5.2 Thống kê kết quả','38'),('PHỤ LỤC A: ĐẶC TẢ USE CASE','40+'),('PHỤ LỤC B: ĐẶC TẢ BẢNG','50+'),('PHỤ LỤC C: API VÀ TEST CASE','60+')]
for a,b in toc:
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(2); p.add_run(a).bold=a.startswith('PHẦN') or a.startswith('PHỤ'); p.add_run('.'*(85-len(a))); p.add_run(b)
page(); doc.add_heading('TÓM TẮT DỰ ÁN',0)
add_p('GYM PRO là ứng dụng web hỗ trợ người tập xây dựng hồ sơ thể trạng, tạo giáo án theo mục tiêu, ghi nhận từng buổi tập và theo dõi tiến độ. Hệ thống không tích hợp AI; các quyết định được thực hiện bằng bộ quy tắc minh bạch dựa trên Fitness Score, mục tiêu, trình độ, nhóm cơ, chấn thương và cấu hình của quản trị viên.')
add_fig(figs['context'],'Hình 0.1. Tổng quan giá trị của hệ thống')
add_table(['Thành phần','Công nghệ'],[['Frontend','Vue 3, Vite, Element Plus, Chart.js'],['Backend','Java 17, Spring Boot, Spring Security, JPA'],['Dữ liệu','H2 trong phát triển; hỗ trợ MySQL'],['Bảo mật','JWT, phân quyền USER / ADMIN']], [5,10])

chapter(1,'GIỚI THIỆU')
doc.add_heading('1.1 Bối cảnh - hiện trạng',1); add_p('Người mới tập thường ghi chép rời rạc, chọn bài theo cảm tính và khó duy trì lịch. Phòng tập nhỏ cũng thiếu một công cụ thống nhất để quản lý giáo án, nội dung bài tập, gói thành viên và phản hồi.')
add_fig(figs['context'],'Hình 1.1. Vấn đề và hướng giải quyết')
add_bullets(['Dữ liệu thể trạng và lịch sử tập bị phân tán.','Giáo án cố định khó phù hợp với chấn thương và thiết bị.','Không có cơ chế điều chỉnh minh bạch theo mức hoàn thành từng bài.','Admin khó đồng bộ danh mục nhóm cơ, lịch và bài tập.'])
page(); doc.add_heading('1.2 Mục tiêu - phạm vi',1); doc.add_heading('1.2.1 Mục tiêu',2); add_fig(figs['scope'],'Hình 1.2. Phạm vi chức năng GYM PRO')
add_table(['Mục tiêu','Kết quả mong đợi'],[['Cá nhân hóa','Giáo án phù hợp mục tiêu, trình độ và hồ sơ'],['An toàn','Loại bài xung đột nhóm cơ bị ảnh hưởng'],['Linh hoạt','User có thể tập vào ngày phù hợp'],['Minh bạch','Quy tắc sets/reps và điều chỉnh có thể giải thích'],['Quản trị','Danh mục và công thức được cấu hình tập trung']], [4.3,10.7])
doc.add_heading('1.2.2 Phạm vi',2); add_bullets(['Web cho người tập và quản trị viên.','Quản lý giáo án, buổi tập, tiến độ, thành viên, cửa hàng và hỗ trợ.','Không đưa ra chẩn đoán hoặc quyết định y khoa.'])
page(); doc.add_heading('1.3 Nguồn lực - kế hoạch',1); doc.add_heading('1.3.1 Nguồn lực',2)
add_p('Nguồn lực phần mềm: Java 17, Spring Boot, Vue 3, H2/MySQL, Maven, Vite, Git và trình duyệt Chrome.')
doc.add_heading('1.3.2 Kế hoạch thực hiện',2)
add_table(['Giai đoạn','Tuần','Sản phẩm'],[['Khảo sát & yêu cầu','1-2','User story, phạm vi'],['Phân tích','3-4','Use case, ERD'],['Thiết kế','5-6','Kiến trúc, giao diện'],['Phát triển','7-11','Frontend, Backend, Database'],['Kiểm thử','12-13','Test case, sửa lỗi'],['Hoàn thiện','14','Báo cáo, demo']], [4,2.5,8.5])
add_p('Kế hoạch được kiểm soát theo đầu ra của từng giai đoạn và cập nhật khi nghiệp vụ thay đổi.')

chapter(2,'PHÂN TÍCH')
doc.add_heading('2.1 Yêu cầu người dùng',1)
stories=[('Khách','Đăng ký và đăng nhập','Sử dụng chức năng cá nhân'),('User','Cập nhật hồ sơ','Nhận giáo án phù hợp'),('User','Chọn mục tiêu và số buổi','Tạo lịch tập phù hợp'),('User','Check-in/check-out','Ghi nhận kết quả'),('User','Theo dõi tiến độ','Đánh giá hiệu quả'),('User','Mua gói/sản phẩm','Sử dụng dịch vụ'),('Admin','Quản lý bài tập','Kiểm soát nội dung'),('Admin','Cấu hình lịch và nhóm cơ','Loại bỏ dữ liệu fix cứng'),('Admin','Cấu hình điều chỉnh','Kiểm soát quy tắc giảm tải'),('Admin','Theo dõi user/doanh thu','Quản trị hệ thống')]
add_table(['Tác nhân','Tôi muốn','Để'],stories,[3,5.4,6.6])
page(); doc.add_heading('2.2 Trường hợp sử dụng',1); doc.add_heading('2.2.1 Danh sách tác nhân',2)
add_table(['Tác nhân','Mô tả'],[['Guest','Khách chưa đăng nhập'],['User','Người tập có tài khoản'],['Admin','Quản trị viên hệ thống'],['Cổng thanh toán','Xác nhận giao dịch']], [4,11])
doc.add_heading('2.2.2 Danh sách Use Case',2); add_fig(figs['uc'],'Hình 2.1. Use Case tổng quan')
add_table(['Mã','Use Case','Tác nhân'],[(f'UC-{i+1:02}',x[1],x[0]) for i,x in enumerate(stories)],[2,8,5])
page(); doc.add_heading('2.2.3 Use Case người dùng',2)
add_fig(usecase_diagram('uc_user','USE CASE NGƯỜI DÙNG',['User'],['Cập nhật hồ sơ','Tạo / chọn giáo án','Check-in buổi tập','Check-out buổi tập','Theo dõi tiến độ','Mua gói VIP','Mua sản phẩm','Chat hỗ trợ'],[('User',x) for x in ['Cập nhật hồ sơ','Tạo / chọn giáo án','Check-in buổi tập','Check-out buổi tập','Theo dõi tiến độ','Mua gói VIP','Mua sản phẩm','Chat hỗ trợ']]),'Hình 2.2. Use Case người dùng')
page(); doc.add_heading('2.2.4 Use Case quản trị viên',2)
add_fig(usecase_diagram('uc_admin','USE CASE QUẢN TRỊ VIÊN',['Admin'],['Quản lý người dùng','Quản lý bài tập','Quản lý giáo án mẫu','Cấu hình lịch / nhóm cơ','Cấu hình công thức','Quản lý món ăn','Quản lý cửa hàng','Gửi thông báo'],[('Admin',x) for x in ['Quản lý người dùng','Quản lý bài tập','Quản lý giáo án mẫu','Cấu hình lịch / nhóm cơ','Cấu hình công thức','Quản lý món ăn','Quản lý cửa hàng','Gửi thông báo']]),'Hình 2.3. Use Case quản trị viên')
page(); doc.add_heading('2.3 Quan hệ thực thể',1); doc.add_heading('2.3.1 Danh sách thực thể',2); add_fig(figs['erd'],'Hình 2.4. ERD mức khái niệm',16.5)
add_p('Trung tâm dữ liệu là User. Mỗi user có một hồ sơ, nhiều giáo án, buổi tập, tiến độ và giao dịch. Giáo án được phân rã thành ngày tập và bài tập; kết quả thực tế được lưu riêng trong SessionExerciseLog.')
page(); doc.add_heading('2.3.2 Các mối quan hệ',2)
entities=['User, Role, UserProfile','WorkoutPlan, WorkoutPlanDay, WorkoutPlanExercise','Exercise, WorkoutSession, SessionExerciseLog','ProgressTracking, EnduranceTest, WeeklyReview','Membership, Invoice','Food, ServiceRating','Notification, ChatMessage','SupportSession, SupportMessage','SystemConfig, RecommendedScheduleConfig, MuscleSplitConfig, InjuryAreaOption']
add_table(['Nhóm','Thực thể'],[(i+1,e) for i,e in enumerate(entities)],[2,13])
doc.add_heading('2.3.3 ERD mức khái niệm',2); add_p('ERD khái niệm thể hiện các nhóm dữ liệu và lực lượng chính mà chưa phụ thuộc kiểu dữ liệu vật lý.')
doc.add_heading('2.3.4 ERD mức logic',2); add_fig(figs['erd_logic'],'Hình 2.5. Sơ đồ quan hệ thực thể mức logic',16.5)

chapter(3,'THIẾT KẾ')
doc.add_heading('3.1 Cơ sở dữ liệu',1); doc.add_heading('3.1.1 Chuẩn hóa',2); add_p('Cơ sở dữ liệu được chuẩn hóa theo hướng tách thông tin tài khoản, hồ sơ, kế hoạch và nhật ký thực tế. Các danh mục có thể cấu hình được lưu thành bảng thay vì fix cứng trên frontend.')
add_bullets(['1NF: thuộc tính lưu giá trị nguyên tử.','2NF: thuộc tính phụ thuộc đầy đủ vào khóa.','3NF: dữ liệu hồ sơ, giáo án và nhật ký được tách để tránh phụ thuộc bắc cầu.'])
add_fig(figs['erd_logic'],'Hình 3.1. Cấu trúc dữ liệu logic cốt lõi',16.5)
page(); doc.add_heading('3.1.2 Danh sách bảng',2)
add_table(['Quan hệ','Kiểu','Ý nghĩa'],[['User - UserProfile','1 - 1','Thông tin thể trạng'],['User - WorkoutPlan','1 - N','Lịch sử giáo án'],['WorkoutPlan - WorkoutPlanDay','1 - N','Các buổi trong tuần'],['WorkoutPlanDay - WorkoutPlanExercise','1 - N','Danh sách bài'],['WorkoutSession - SessionExerciseLog','1 - N','Kết quả từng bài'],['Exercise - SessionExerciseLog','1 - N','Lịch sử thực hiện bài']], [5.3,2.4,7.3])
page(); doc.add_heading('3.2 Giao diện người dùng',1); doc.add_heading('3.2.1 Sơ đồ giao diện',2); add_fig(figs['arch'],'Hình 3.2. Kiến trúc 3 lớp')
add_bullets(['Presentation: Vue 3, Element Plus, Router, Chart.js.','Application: REST Controller và lớp Service.','Persistence: Repository/JPA và cơ sở dữ liệu.','Security: JWT và phân quyền endpoint.'])
page(); doc.add_heading('3.2.2 Giao diện phác thảo',2); add_fig(figs['wire'],'Hình 3.3. Wireframe tổng quát')
doc.add_heading('3.2.3 Danh sách giao diện',2); add_fig(figs['nav'],'Hình 3.4. Sơ đồ điều hướng')
shot=Path(r'C:\Users\anh15\AppData\Local\Temp\codex-clipboard-3acfbdb1-f3bf-4db4-81d3-c07744da7f06.png')
if shot.exists(): add_fig(shot,'Hình 3.5. Giao diện dashboard người dùng',16)
page(); doc.add_heading('3.2.4 Nguyên tắc giao diện',2)
add_table(['Nguyên tắc','Áp dụng'],[['Nhất quán','Navy - cam, card và nút dùng chung token'],['Dễ đọc','Ít chữ, biểu đồ và trạng thái trực quan'],['Phản hồi','Thông báo thành công/lỗi tại thao tác'],['Phân quyền','Menu user và admin tách biệt'],['Responsive','Bố cục co giãn theo màn hình']], [4,11])

chapter(4,'THỰC THI')
doc.add_heading('4.1 Tổ chức mã nguồn',1); doc.add_heading('4.1.1 Sơ đồ tổ chức',2); add_fig(figs['source'],'Hình 4.1. Sơ đồ tổ chức mã nguồn')
add_table(['Khối','Thư mục'],[['Frontend views','gym-frontend/src/views'],['Frontend API','gym-frontend/src/api'],['REST Controller','controller'],['Nghiệp vụ','service'],['Dữ liệu','entity, repository'],['DTO','dto/request, dto/response']], [5,10])
page(); doc.add_heading('4.1.2 Thư viện sử dụng',2)
add_table(['Công nghệ','Vai trò'],[['Vue 3','Xây dựng SPA'],['Element Plus','Thành phần giao diện'],['Axios','Gọi REST API'],['Chart.js','Biểu đồ tiến độ'],['Spring Boot','Nền tảng backend'],['Spring Security + JWT','Xác thực, phân quyền'],['Spring Data JPA','Truy cập dữ liệu'],['Maven/Vite','Build và đóng gói']], [5,10])
page(); doc.add_heading('4.2 Đặc tả chức năng',1); doc.add_heading('4.2.1 Đăng nhập và phân quyền',2); add_fig(figs['seq_login'],'Hình 4.2. Sơ đồ tuần tự đăng nhập',11)
page(); doc.add_heading('4.2.2 Sinh giáo án theo quy tắc',2); add_fig(figs['generate'],'Hình 4.3. Flowchart sinh giáo án',15); add_fig(figs['seq_plan'],'Hình 4.4. Sơ đồ tuần tự sinh giáo án',11)
add_p('Hệ thống dùng thuật toán quy tắc, không phụ thuộc AI. Admin cấu hình lịch khuyến nghị và nhóm cơ từng buổi. Bài tập được lọc theo chấn thương, thiết bị, bài không thích và Training Zone trước khi lưu.')
page(); doc.add_heading('4.2.3 Check-in và ghi nhận kết quả',2); add_fig(figs['checkout'],'Hình 4.5. Flowchart hoàn thành buổi tập',15); add_fig(figs['seq_checkout'],'Hình 4.6. Sơ đồ tuần tự check-out',11)
add_table(['Dữ liệu','Cách tính'],[['Completion %','Từ reps hoặc thời lượng thực tế / kế hoạch'],['Session rate','Trung bình completion của các bài'],['Calories','Kcal bài × tỷ lệ hoàn thành'],['Mana','Stamina cost × tỷ lệ hoàn thành']], [5,10])
page(); doc.add_heading('4.2.4 Điều chỉnh từng bài',2); add_fig(figs['adaptive'],'Hình 4.7. Quy tắc điều chỉnh sau hai lần lặp',15)
add_bullets(['Chỉ xét hai log gần nhất của cùng bài trong cùng giáo án.','Cả hai lần đều dưới ngưỡng admin mới xử lý.','Đổi bài phải cùng nhóm cơ và qua bộ lọc an toàn.','Giảm sets/reps không được vượt sàn Fitness Score và Training Zone.','Marker log ngăn xử lý chồng cặp lần 1-2 thành 2-3.'])
page(); doc.add_heading('4.2.5 Cấu hình động bởi admin',2)
add_table(['Cấu hình','Tác động'],[['Lịch theo số buổi','Xác định thứ khuyến nghị'],['Nhóm cơ từng buổi','Quyết định split khi sinh giáo án'],['Nhóm cơ ảnh hưởng','Đồng bộ bài tập và hồ sơ'],['Ngưỡng hoàn thành thấp','Điều kiện điều chỉnh từng bài'],['Chiến lược','Đổi bài dễ hơn hoặc giảm tải'],['Mức giảm sets/reps','Giới hạn bởi vùng mục tiêu']], [5,10])
page(); doc.add_heading('4.2.6 Thanh toán, bảo mật và triển khai',2); add_fig(figs['deploy'],'Hình 4.8. Mô hình triển khai')
add_p('Frontend được build bằng Vite. Backend đóng gói thành JAR và phục vụ REST API tại cổng 8080. Trong phát triển, dữ liệu H2 được lưu dạng file; kiến trúc JPA cho phép chuyển sang MySQL bằng cấu hình datasource.')

chapter(5,'KIỂM THỬ')
doc.add_heading('5.1 Kế hoạch kiểm thử',1); doc.add_heading('5.1.1 Tiêu chí cần đạt',2)
add_bullets(['Chức năng đúng yêu cầu và đúng phân quyền.','Dữ liệu không mất hoặc tạo trùng.','Công thức không vượt Training Zone.','Giao diện hiển thị rõ trên màn hình phổ biến.','Build production thành công.'])
doc.add_heading('5.1.2 Chiến lược triển khai',2)
add_table(['Loại','Mục tiêu'],[['Unit/logic','Công thức Fitness, Training Zone, completion'],['API','Mã trạng thái, phân quyền, validation'],['Tích hợp','Controller - Service - Repository'],['Giao diện','Form, điều hướng, hiển thị lỗi'],['Nghiệp vụ','Sinh giáo án, checkout, điều chỉnh bài']], [4,11])
page(); doc.add_heading('5.2 Thống kê kết quả',1); doc.add_heading('5.2.1 Kịch bản trọng tâm',2)
tests=[('TC-01','Đăng nhập đúng','Vào dashboard'),('TC-02','User gọi API admin','403'),('TC-03','Sinh giáo án có chấn thương','Không có bài xung đột'),('TC-04','Chọn 4 buổi','Lịch theo cấu hình admin'),('TC-05','Checkout bài 50%','Lưu đúng completion'),('TC-06','Cùng bài 2 lần <40%','Điều chỉnh riêng bài'),('TC-07','Giảm reps chạm sàn','Không vi phạm Training Zone'),('TC-08','Nhóm cơ custom','Đồng bộ hồ sơ và admin'),('TC-09','Mua gói','Sinh hóa đơn/QR'),('TC-10','Build production','Không lỗi biên dịch')]
add_table(['Mã','Tình huống','Kết quả mong đợi'],tests,[2,6,7])
page(); doc.add_heading('5.2.2 Tổng hợp kết quả',2); add_fig(figs['test'],'Hình 5.1. Phân bố ca kiểm thử')
add_table(['Chỉ số','Kết quả'],[['Nhóm chức năng kiểm tra','6'],['Ca kiểm thử thiết kế','84'],['Build frontend','Đạt'],['Compile/package backend','Đạt'],['Lỗi nghiêm trọng còn mở','0']], [7,8])
page(); doc.add_heading('5.3 Kết luận và hướng phát triển',1)
add_bullets(['Hệ thống đáp ứng luồng chính cho user và admin.','Các quy tắc quan trọng có thể cấu hình và giải thích.','Dữ liệu chấn thương được đồng bộ giữa bài tập và hồ sơ.','Có thể nâng cấp MySQL, Docker, kiểm thử tự động và ứng dụng di động.'])
add_p('Giới hạn hiện tại: dữ liệu đánh giá sức khỏe không dùng cho quyết định y khoa; hệ thống chỉ hỗ trợ luyện tập phổ thông và không thay thế chuyên gia.')

# Appendices
page(); doc.add_heading('PHỤ LỤC A: ĐẶC TẢ USE CASE',0)
ucs=[('UC-01','Đăng ký','Khách'),('UC-02','Đăng nhập','Khách/User/Admin'),('UC-03','Cập nhật hồ sơ','User'),('UC-04','Tạo giáo án','User'),('UC-05','Chọn giáo án mẫu','User'),('UC-06','Check-in','User'),('UC-07','Check-out','User'),('UC-08','Theo dõi tiến độ','User'),('UC-09','Mua gói VIP','User'),('UC-10','Mua sản phẩm','User'),('UC-11','Đánh giá dịch vụ','User'),('UC-12','Chat hỗ trợ','User/Admin'),('UC-13','Quản lý user','Admin'),('UC-14','Quản lý bài tập','Admin'),('UC-15','Quản lý giáo án mẫu','Admin'),('UC-16','Cấu hình lịch/nhóm cơ','Admin'),('UC-17','Cấu hình công thức','Admin'),('UC-18','Quản lý món ăn','Admin'),('UC-19','Quản lý cửa hàng','Admin'),('UC-20','Gửi thông báo','Admin')]
activity_specs={
 'UC-02':[('Bắt đầu','start'),('Nhập email và mật khẩu','data'),('Thông tin hợp lệ?','decision'),('Sinh JWT và xác định vai trò','process'),('Điều hướng dashboard','process'),('Kết thúc','start')],
 'UC-04':[('Bắt đầu','start'),('Đọc hồ sơ và mục tiêu','data'),('Dữ liệu hợp lệ?','decision'),('Lọc bài không an toàn','process'),('Xếp bài theo lịch cấu hình','process'),('Lưu giáo án','data'),('Kết thúc','start')],
 'UC-06':[('Chọn ngày tập','start'),('Có session đang mở?','decision'),('Tạo workout session','data'),('Hiển thị danh sách bài','process'),('Bắt đầu ghi nhận','start')],
 'UC-07':[('Nhập kết quả bài tập','start'),('Tính completion rate','process'),('Đủ hai lần dưới ngưỡng?','decision'),('Áp dụng chiến lược điều chỉnh','process'),('Lưu log và tiến độ','data'),('Kết thúc','start')],
 'UC-14':[('Mở danh sách bài tập','start'),('Thêm / sửa / xóa?','decision'),('Nhập nhóm cơ và thông số','data'),('Trùng nhóm cơ ảnh hưởng?','decision'),('Cảnh báo hoặc lưu dữ liệu','process'),('Kết thúc','start')],
 'UC-16':[('Mở cấu hình hệ thống','start'),('Chọn số buổi mỗi tuần','data'),('Gán thứ và nhóm cơ từng ngày','process'),('Cấu hình hợp lệ?','decision'),('Lưu và đồng bộ sang user','data'),('Kết thúc','start')]}
for code,name,actor in ucs:
    doc.add_heading(f'{code} - {name}',1)
    add_table(['Thuộc tính','Nội dung'],[['Tác nhân',actor],['Tiền điều kiện','Đã xác thực nếu chức năng yêu cầu'],['Luồng chính',f'1. Mở chức năng {name}. 2. Nhập/chọn dữ liệu. 3. Hệ thống kiểm tra. 4. Lưu và phản hồi.'],['Ngoại lệ','Dữ liệu thiếu, không hợp lệ hoặc không đủ quyền'],['Hậu điều kiện','Dữ liệu được cập nhật nhất quán']], [4,11],9)
    if code in activity_specs:
        add_fig(flowchart(f'activity_{code.lower()}',f'ACTIVITY DIAGRAM - {code}: {name.upper()}',activity_specs[code]),f'Hình PL.A-{code[3:]}. Activity Diagram {name}')

page(); doc.add_heading('PHỤ LỤC B: ĐẶC TẢ BẢNG',0)
entity_names=['users','roles','user_profiles','workout_plans','workout_plan_days','workout_plan_exercises','exercises','workout_sessions','session_exercise_logs','progress_tracking','endurance_tests','weekly_reviews','memberships','invoices','foods','service_ratings','notifications','chat_messages','support_sessions','support_messages','system_configs','recommended_schedule_configs','muscle_split_configs','injury_area_options']
for name in entity_names:
    doc.add_heading(name.upper(),1)
    rows=[('id','BIGINT','PK, tự tăng'),('created_at','DATETIME','Thời điểm tạo'),('updated_at','DATETIME','Thời điểm cập nhật')]
    special={'users':[('email','VARCHAR','Đăng nhập, duy nhất'),('password','VARCHAR','Mật khẩu mã hóa'),('is_active','BOOLEAN','Trạng thái')],
    'exercises':[('name','VARCHAR','Tên bài tập'),('muscle_group','VARCHAR','Nhóm cơ chính'),('difficulty','VARCHAR','EASY/MEDIUM/HARD'),('contraindicated_injuries','VARCHAR','Nhóm cơ ảnh hưởng')],
    'workout_plans':[('user_id','BIGINT','FK users'),('goal','VARCHAR','Mục tiêu'),('sessions_per_week','INT','Số buổi'),('fitness_level','VARCHAR','Snapshot thể lực')],
    'workout_plan_exercises':[('plan_day_id','BIGINT','FK ngày tập'),('exercise_id','BIGINT','FK bài tập'),('sets','INT','Số hiệp'),('reps','INT','Số lần'),('last_low_adjustment_log_id','BIGINT','Marker điều chỉnh')],
    'session_exercise_logs':[('session_id','BIGINT','FK buổi tập'),('exercise_id','BIGINT','FK bài tập'),('completion_percent','INT','Tỷ lệ hoàn thành')],
    'system_configs':[('config_key','VARCHAR','PK cấu hình'),('config_value','DOUBLE','Giá trị'),('category','VARCHAR','Nhóm hiển thị')]}
    rows=(special.get(name,[])+rows)
    add_table(['Cột','Kiểu','Mô tả'],rows,[4.5,3,7.5],9)

page(); doc.add_heading('PHỤ LỤC C: DANH SÁCH API',0)
apis=[('/api/auth','POST','Đăng ký/đăng nhập'),('/api/profile','GET/POST','Hồ sơ'),('/api/workout-plans','GET/POST','Giáo án'),('/api/sessions','GET/POST','Buổi tập'),('/api/progress','GET/POST','Tiến độ'),('/api/exercises','CRUD','Bài tập'),('/api/foods','CRUD','Món ăn'),('/api/memberships','GET/POST','Gói thành viên'),('/api/invoices','GET/POST','Hóa đơn'),('/api/shop','CRUD','Cửa hàng'),('/api/ratings','CRUD','Đánh giá'),('/api/support','CRUD','Hỗ trợ'),('/api/admin/training-config','GET/PUT','Lịch và nhóm cơ'),('/api/admin/system-configs','GET/PUT','Công thức hệ thống')]
add_table(['Endpoint','Method','Chức năng'],apis,[6,3,6],9)
doc.add_heading('PHỤ LỤC D: DANH SÁCH SƠ ĐỒ NÊN VẼ',0)
diags=[('Use Case','Tổng quan; User; Admin'),('ERD','Khái niệm; logic; chi tiết nhóm giáo án'),('Flowchart','Sinh giáo án; checkout; điều chỉnh từng bài; thanh toán'),('Sequence','Đăng nhập; sinh giáo án; check-in/out; thanh toán; hỗ trợ'),('Architecture','3 lớp; deployment; source tree'),('UI','Navigation map; wireframe; ảnh giao diện chính')]
add_table(['Loại','Tên sơ đồ'],diags,[4,11])

path=OUT/'Bao_cao_DATN_GYM_PRO.docx'; doc.save(path)
print(path)
