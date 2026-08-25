from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "ON_TAP_CHUC_NANG_LOI_VA_DEMO_20_PHUT.md"
OUTPUT = ROOT / "ON_TAP_CHUC_NANG_LOI_VA_DEMO_20_PHUT.docx"

BLUE = "2E74B5"
DARK = "163A5F"
LIGHT = "E8EEF5"
PALE = "F4F7FA"
GOLD = "C48A22"
GRAY = "5F6B76"
WHITE = "FFFFFF"
INK = "1F2933"


def set_font(run, size=None, bold=None, italic=None, color=None, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths):
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[min(index, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            set_cell_margins(cell)


def set_keep(paragraph, next_=False):
    p_pr = paragraph._p.get_or_add_pPr()
    keep = OxmlElement("w:keepNext" if next_ else "w:keepLines")
    p_pr.append(keep)


def add_inline(paragraph, text, default_bold=False, color=INK):
    # Render Markdown bold and inline code without leaving markup characters.
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_font(run, bold=True, color=color)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_font(run, size=9.2, bold=True, color=DARK, name="Consolas")
        else:
            run = paragraph.add_run(part)
            set_font(run, bold=default_bold, color=color)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.16
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 9),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 11.5, DARK, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.16


def add_header_footer(section):
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)
    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    run = hp.add_run("GYMIFY  /  TÀI LIỆU ÔN TẬP BẢO VỆ")
    set_font(run, size=8.5, bold=True, color=GRAY)
    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(0)
    r = fp.add_run("Gymify • Kịch bản demo 20 phút")
    set_font(r, size=8.5, color=GRAY)


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(34)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("ĐỒ ÁN TỐT NGHIỆP")
    set_font(r, size=10, bold=True, color=GOLD)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("GYMIFY")
    set_font(r, size=30, bold=True, color=DARK)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Tài liệu ôn chức năng lõi\nvà kịch bản demo 20 phút")
    set_font(r, size=20, bold=True, color=BLUE)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(24)
    r = p.add_run("Xây dựng website lập lịch tập và cá nhân hóa giáo án luyện tập cho phòng gym Gymify")
    set_font(r, size=12, italic=True, color=GRAY)

    table = doc.add_table(rows=1, cols=4)
    labels = [("03 phút", "Bài toán"), ("08 phút", "Cá nhân hóa"), ("05 phút", "Tập & tiến độ"), ("04 phút", "CSDL & kết luận")]
    for i, (time, label) in enumerate(labels):
        cell = table.cell(0, i)
        shade(cell, LIGHT if i % 2 == 0 else PALE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(time + "\n")
        set_font(r, size=11, bold=True, color=DARK)
        r = p.add_run(label)
        set_font(r, size=8.5, color=GRAY)
    set_table_geometry(table, [2340] * 4)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(28)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("THÔNG ĐIỆP PHẢI CHỨNG MINH")
    set_font(r, size=9.5, bold=True, color=GOLD)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.18)
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("Hai hồ sơ khác nhau tạo ra lịch, bài tập và tải tập khác nhau; mọi khác biệt đều truy ngược được về dữ liệu đầu vào và quy tắc nghiệp vụ.")
    set_font(r, size=13, bold=True, color=DARK)
    doc.add_page_break()


def add_callout(doc, text):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    shade(cell, PALE)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_inline(p, text, color=DARK)
    set_table_geometry(table, [9360])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_markdown_table(doc, rows):
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = table.cell(r_idx, c_idx)
            value = row[c_idx] if c_idx < len(row) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            add_inline(p, value, default_bold=(r_idx == 0), color=(WHITE if r_idx == 0 else INK))
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if r_idx == 0:
                shade(cell, DARK)
            elif r_idx % 2 == 0:
                shade(cell, PALE)
    set_repeat_header(table.rows[0])
    # Content-aware patterns for the 2-, 3-, and 4-column tables in this guide.
    width_map = {
        2: [2700, 6660],
        3: [2300, 2530, 4530],
        4: [1900, 2050, 2050, 3360],
    }
    set_table_geometry(table, width_map.get(cols, [9360 // cols] * cols))
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def is_separator(line):
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return cells and all(re.fullmatch(r":?-{3,}:?", c) for c in cells)


def parse_table(lines, start):
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        if not is_separator(lines[i]):
            rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
        i += 1
    return rows, i


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    configure_styles(doc)
    add_header_footer(section)
    add_cover(doc)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    in_code = False
    code_lines = []
    i = 1  # Skip Markdown H1; cover replaces it.
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        if line.startswith("```"):
            if in_code:
                add_callout(doc, "\n".join(code_lines))
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(raw)
            i += 1
            continue
        if not line:
            i += 1
            continue
        if line.startswith("|") and i + 1 < len(lines) and is_separator(lines[i + 1]):
            rows, i = parse_table(lines, i)
            add_markdown_table(doc, rows)
            continue
        heading = re.match(r"^(#{2,4})\s+(.*)$", line)
        if heading:
            level = min(len(heading.group(1)) - 1, 3)
            p = doc.add_paragraph(style=f"Heading {level}")
            add_inline(p, heading.group(2), default_bold=True, color=BLUE if level < 3 else DARK)
            set_keep(p, next_=True)
            i += 1
            continue
        if line.startswith(">"):
            add_callout(doc, line.lstrip("> "))
            i += 1
            continue
        numbered = re.match(r"^\d+\.\s+(.*)$", line)
        if numbered:
            p = doc.add_paragraph(style="List Number")
            add_inline(p, numbered.group(1))
            i += 1
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, line[2:])
            i += 1
            continue
        p = doc.add_paragraph()
        add_inline(p, line)
        i += 1

    core = doc.core_properties
    core.title = "Gymify - Tài liệu ôn chức năng lõi và kịch bản demo 20 phút"
    core.subject = "Tài liệu bảo vệ đồ án Gymify"
    core.author = "Gymify"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
