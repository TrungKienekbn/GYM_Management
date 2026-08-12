from pathlib import Path
import re
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(r"C:\Users\anh15\GYM_Management")
SOURCE = Path(sys.argv[1]) if len(sys.argv) > 1 else ROOT / "TAI_LIEU_DEMO_30_PHUT.md"
OUTPUT = Path(sys.argv[2]) if len(sys.argv) > 2 else ROOT / "TAI_LIEU_DEMO_30_PHUT.docx"
SPEC_MODE = SOURCE.stem.startswith("DAC_TA_CHUC_NANG")

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "172033"
ORANGE = "D97706"
LIGHT_BLUE = "E8EEF5"
LIGHT_ORANGE = "FFF4E5"
GRAY = "667085"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths_dxa):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_font(run, name="Calibri", size=None, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_inline(paragraph, text, base_size=11, base_color=None):
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_font(run, "Consolas", 9.5, color=NAVY)
            run._element.get_or_add_rPr().append(OxmlElement("w:shd"))
            run._element.rPr[-1].set(qn("w:fill"), "F2F4F7")
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_font(run, size=base_size, bold=True, color=base_color)
        else:
            run = paragraph.add_run(part)
            set_font(run, size=base_size, color=base_color)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Trang ")
    set_font(run, size=9, color=GRAY)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.18

    header = section.header
    hp = header.paragraphs[0]
    hp.text = "GYMPRO  |  ĐẶC TẢ CHỨC NĂNG" if SPEC_MODE else "GYMPRO  |  KỊCH BẢN DEMO 30 PHÚT"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(hp.runs[0], size=8.5, bold=True, color=GRAY)
    footer = section.footer
    add_page_number(footer.paragraphs[0])


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(115)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("GYMPRO")
    set_font(r, size=34, bold=True, color=NAVY)
    r2 = p.add_run("  FUNCTIONAL SPECIFICATION" if SPEC_MODE else "  DEMO GUIDE")
    set_font(r2, size=13, bold=True, color=ORANGE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("ĐẶC TẢ CHỨC NĂNG HỆ THỐNG" if SPEC_MODE else "KỊCH BẢN THUYẾT TRÌNH 30 PHÚT")
    set_font(r, size=25, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Phạm vi - tác nhân - nghiệp vụ - quy tắc - ngoại lệ" if SPEC_MODE else "So sánh hai hồ sơ - hai mục tiêu - hai giáo án")
    set_font(r, size=14, italic=True, color=DARK_BLUE)

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    labels = (("33", "chức năng"), ("4", "nhóm tác nhân"), ("2", "vai trò chính"), ("7", "nhóm chất lượng")) if SPEC_MODE else (("4", "người trình bày"), ("30", "phút"), ("2", "hồ sơ đối chứng"), ("1", "luồng nghiệp vụ"))
    for idx, (value, label) in enumerate(labels):
        cell = table.cell(0, idx)
        set_cell_shading(cell, LIGHT_ORANGE if idx % 2 == 0 else LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(value + "\n")
        set_font(r, size=18, bold=True, color=ORANGE if idx % 2 == 0 else BLUE)
        r = p.add_run(label)
        set_font(r, size=8.5, bold=True, color=NAVY)
    set_table_widths(table, [2340, 2340, 2340, 2340])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    r = p.add_run("Tài liệu dùng cho báo cáo, kiểm thử và thuyết trình" if SPEC_MODE else "Tài liệu dùng khi demo và trả lời phản biện")
    set_font(r, size=10.5, color=GRAY)
    doc.add_page_break()


def add_toc(doc, headings):
    p = doc.add_paragraph("MỤC LỤC NỘI DUNG", style="Heading 1")
    p.paragraph_format.space_before = Pt(0)
    for level, title in headings:
        if level > 2:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25 if level == 2 else 0)
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(title)
        set_font(r, size=10.5 if level == 2 else 11, bold=(level == 1), color=DARK_BLUE if level == 1 else GRAY)
    doc.add_page_break()


def parse_table(lines, index):
    rows = []
    while index < len(lines) and lines[index].strip().startswith("|"):
        cells = [c.strip() for c in lines[index].strip().strip("|").split("|")]
        rows.append(cells)
        index += 1
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in rows[1]):
        rows.pop(1)
    return rows, index


def table_widths(count):
    if count == 2:
        return [2700, 6660]
    if count == 3:
        return [2200, 3580, 3580]
    if count == 4:
        return [1450, 2100, 2810, 3000]
    return [9360 // count] * count


def add_markdown_table(doc, rows):
    if not rows:
        return
    cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for ridx, row in enumerate(rows):
        for cidx in range(cols):
            cell = table.cell(ridx, cidx)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_inline(p, row[cidx] if cidx < len(row) else "", base_size=8.5 if cols >= 4 else 9.3)
            if ridx == 0:
                set_cell_shading(cell, LIGHT_BLUE)
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(NAVY)
            elif ridx % 2 == 0:
                set_cell_shading(cell, "F8FAFC")
    set_repeat_table_header(table.rows[0])
    set_table_widths(table, table_widths(cols))
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def add_quote(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_ORANGE)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_font(r, size=10.2, italic=True, color=NAVY)
    set_table_widths(table, [9360])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_code_block(doc, content):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F2F4F7")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for i, line in enumerate(content):
        r = p.add_run(("" if i == 0 else "\n") + line)
        set_font(r, "Consolas", 9.2, color=NAVY)
    set_table_widths(table, [9360])


def build_doc():
    text = SOURCE.read_text(encoding="utf-8")
    lines = text.splitlines()
    headings = []
    for line in lines:
        m = re.match(r"^(#{1,3})\s+(.+)$", line)
        if m and not m.group(2).startswith("KỊCH BẢN DEMO"):
            headings.append((len(m.group(1)), m.group(2)))

    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_toc(doc, headings)

    in_code = False
    code_lines = []
    i = 1  # Skip Markdown document title; cover already provides it.
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        if line.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(raw)
            i += 1
            continue
        if not line:
            i += 1
            continue
        if line.startswith("|"):
            rows, i = parse_table(lines, i)
            add_markdown_table(doc, rows)
            continue
        m = re.match(r"^(#{1,3})\s+(.+)$", line)
        if m:
            level = len(m.group(1))
            title = m.group(2)
            p = doc.add_paragraph(title, style=f"Heading {level}")
            if level == 1 and title.startswith(("3.1.", "4.", "5.", "6.", "7.")):
                p.paragraph_format.page_break_before = True
            i += 1
            continue
        if line.startswith(">"):
            quote = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                quote.append(lines[i].strip()[1:].strip())
                i += 1
            add_quote(doc, "\n".join(quote))
            continue
        m = re.match(r"^[-*]\s+(.+)$", line)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, m.group(1))
            i += 1
            continue
        m = re.match(r"^\d+\.\s+(.+)$", line)
        if m:
            p = doc.add_paragraph(style="List Number")
            add_inline(p, m.group(1))
            i += 1
            continue
        p = doc.add_paragraph()
        add_inline(p, line)
        i += 1

    props = doc.core_properties
    props.title = "Đặc tả chức năng hệ thống GymPro" if SPEC_MODE else "Kịch bản demo GymPro 30 phút"
    props.subject = "Đặc tả tác nhân, luồng nghiệp vụ, ngoại lệ và yêu cầu phi chức năng" if SPEC_MODE else "Phân công 4 người và so sánh hai giáo án cá nhân hóa"
    props.author = "Nhóm GymPro"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_doc()
