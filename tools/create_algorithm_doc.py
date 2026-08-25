from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

OUT = Path(r"C:\Users\anh15\GYM_Management\output\TAI_LIEU_THUAT_TOAN_GYMPRO.docx")
NAVY = "16324F"; BLUE = "2E74B5"; LIGHT = "E8EEF5"; PALE = "F4F6F9"; GRAY = "5B6573"; WHITE = "FFFFFF"; BLACK = "111111"

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = tcPr.find(qn('w:shd'))
    if shd is None: shd = OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'), fill)

def margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc.get_or_add_tcPr(); node = tc.first_child_found_in('w:tcMar')
    if node is None: node = OxmlElement('w:tcMar'); tc.append(node)
    for tag, val in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        x=OxmlElement('w:'+tag); x.set(qn('w:w'),str(val)); x.set(qn('w:type'),'dxa'); node.append(x)

def borders(table, color="BCC6D1", size="5"):
    pr=table._tbl.tblPr; b=pr.first_child_found_in('w:tblBorders')
    if b is None: b=OxmlElement('w:tblBorders'); pr.append(b)
    for edge in ('top','left','bottom','right','insideH','insideV'):
        e=OxmlElement('w:'+edge); e.set(qn('w:val'),'single'); e.set(qn('w:sz'),size); e.set(qn('w:color'),color); b.append(e)

def repeat_header(row):
    trPr=row._tr.get_or_add_trPr(); el=OxmlElement('w:tblHeader'); el.set(qn('w:val'),'true'); trPr.append(el)

def cant_split(row):
    trPr=row._tr.get_or_add_trPr(); trPr.append(OxmlElement('w:cantSplit'))

def set_cell_text(cell, text, bold=False, color=BLACK, size=9.5, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text=''; p=cell.paragraphs[0]; p.alignment=align; p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.08
    r=p.add_run(str(text)); r.bold=bold; r.font.name='Calibri'; r.font.size=Pt(size); r.font.color.rgb=RGBColor.from_string(color)
    r._element.get_or_add_rPr().rFonts.set(qn('w:ascii'),'Calibri'); r._element.rPr.rFonts.set(qn('w:hAnsi'),'Calibri')
    cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; margins(cell)

def table(doc, headers, rows, widths=None):
    t=doc.add_table(rows=1, cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False; borders(t)
    for i,h in enumerate(headers): shade(t.rows[0].cells[i], LIGHT); set_cell_text(t.rows[0].cells[i],h,True,NAVY,9.5,WD_ALIGN_PARAGRAPH.CENTER)
    repeat_header(t.rows[0])
    for row in rows:
        cells=t.add_row().cells; cant_split(t.rows[-1])
        for i,v in enumerate(row): set_cell_text(cells[i],v)
    if widths:
        for row in t.rows:
            for i,w in enumerate(widths): row.cells[i].width=Inches(w)
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(1)
    return t

def add_field(paragraph, field):
    r=paragraph.add_run(); begin=OxmlElement('w:fldChar'); begin.set(qn('w:fldCharType'),'begin'); instr=OxmlElement('w:instrText'); instr.set(qn('xml:space'),'preserve'); instr.text=field; sep=OxmlElement('w:fldChar'); sep.set(qn('w:fldCharType'),'separate'); txt=OxmlElement('w:t'); txt.text='1'; end=OxmlElement('w:fldChar'); end.set(qn('w:fldCharType'),'end')
    for x in (begin,instr,sep,txt,end): r._r.append(x)

doc=Document(); sec=doc.sections[0]
sec.page_width=Inches(8.5); sec.page_height=Inches(11); sec.top_margin=Inches(0.85); sec.bottom_margin=Inches(0.75); sec.left_margin=Inches(0.9); sec.right_margin=Inches(0.9); sec.header_distance=Inches(.35); sec.footer_distance=Inches(.35)

styles=doc.styles
normal=styles['Normal']; normal.font.name='Calibri'; normal.font.size=Pt(11); normal.font.color.rgb=RGBColor.from_string(BLACK); normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.15
for name,size,before,after,color in [('Title',27,0,8,NAVY),('Subtitle',13,0,16,GRAY),('Heading 1',16,16,8,BLUE),('Heading 2',13,12,6,BLUE),('Heading 3',11.5,8,4,NAVY)]:
    s=styles[name]; s.font.name='Calibri'; s.font.size=Pt(size); s.font.color.rgb=RGBColor.from_string(color); s.font.bold=name!='Subtitle'; s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True

for base,new in [('Normal','CodeBlock'),('Normal','Callout')]:
    if new not in styles: styles.add_style(new,WD_STYLE_TYPE.PARAGRAPH)
code=styles['CodeBlock']; code.font.name='Consolas'; code.font.size=Pt(9); code.paragraph_format.left_indent=Inches(.25); code.paragraph_format.right_indent=Inches(.15); code.paragraph_format.space_before=Pt(4); code.paragraph_format.space_after=Pt(8); code.paragraph_format.line_spacing=1.0
call=styles['Callout']; call.font.name='Calibri'; call.font.size=Pt(10.5); call.font.color.rgb=RGBColor.from_string(NAVY); call.paragraph_format.left_indent=Inches(.18); call.paragraph_format.right_indent=Inches(.18); call.paragraph_format.space_before=Pt(5); call.paragraph_format.space_after=Pt(8)

header=sec.header.paragraphs[0]; header.text='GYMPRO  |  TÀI LIỆU THUẬT TOÁN'; header.style=styles['Normal']; header.runs[0].font.size=Pt(8); header.runs[0].font.color.rgb=RGBColor.from_string(GRAY)
footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.RIGHT; r=footer.add_run('Trang '); r.font.size=Pt(9); r.font.color.rgb=RGBColor.from_string(GRAY); add_field(footer,'PAGE')

doc.add_paragraph('GYMPRO',style='Subtitle').alignment=WD_ALIGN_PARAGRAPH.CENTER
p=doc.add_paragraph('TÀI LIỆU THUẬT TOÁN\nSINH GIÁO ÁN TẬP LUYỆN CÁ NHÂN HÓA',style='Title'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
p=doc.add_paragraph('Mô tả thuật toán rule-based, công thức, giả mã và ví dụ kiểm chứng theo mã nguồn hiện tại',style='Subtitle'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('\n')
table(doc,['Thông tin','Nội dung'],[
 ('Hệ thống','GymPro - Website quản lý phòng tập và hỗ trợ luyện tập'),('Phạm vi','Thuật toán tạo và cá nhân hóa giáo án'),('Công nghệ','Java 17, Spring Boot, Spring Data JPA'),('Phiên bản tài liệu','1.0 - 15/08/2026'),('Nguồn đối chiếu','FitnessCalculator, MuscleGroupSplitPlanner, TrainingZone, WorkoutPlanService')],[1.55,4.95])
p=doc.add_paragraph('Mục đích sử dụng',style='Heading 2');
doc.add_paragraph('Tài liệu có thể đưa vào chương “Phân tích và thiết kế thuật toán” của báo cáo đồ án. Nội dung tập trung vào khả năng giải thích, tính xác định và các ràng buộc an toàn của hệ thống; không khẳng định đây là chẩn đoán hoặc chỉ định y khoa.',style='Callout')
doc.add_page_break()

doc.add_heading('1. Tổng quan thuật toán',1)
doc.add_paragraph('GymPro sử dụng thuật toán dựa trên luật (rule-based). Cùng một hồ sơ, cấu hình và dữ liệu bài tập sẽ tạo ra cùng một kết quả. Chuỗi xử lý gồm sáu giai đoạn:')
for x in ['Chuẩn hóa hồ sơ và tính Fitness Score (FS).','Phân loại mức thể lực và thể trạng cơ thể.','Xác định số buổi, lịch tập và nhóm cơ của từng buổi.','Phân bổ quota bài tập theo nhóm cơ bằng Largest Remainder Method.','Lọc cứng ứng viên theo an toàn, thiết bị và sở thích; sau đó xếp hạng theo mục tiêu.','Cá nhân hóa sets, reps, thời lượng, nghỉ và mức tạ; giới hạn tổng số bài theo thời lượng buổi.']:
    doc.add_paragraph(x,style='List Number')
doc.add_heading('1.1. Đầu vào và đầu ra',2)
table(doc,['Nhóm','Dữ liệu chính','Vai trò'],[
 ('Hồ sơ','Tuổi, giới tính, chiều cao, cân nặng, BMI, % mỡ','Tính FS và BodyType'),('Mục tiêu','Tăng cơ, giảm cân, sức bền, duy trì','Chọn split, score, training zone và thời gian nghỉ'),('Khả năng','Trình độ, kinh nghiệm, ngày rảnh, thời lượng','Giới hạn tải và quy mô giáo án'),('Ràng buộc','Chấn thương, thiết bị, địa điểm, bài không thích','Loại ứng viên không phù hợp'),('Danh mục bài','Nhóm cơ, độ khó, score, chống chỉ định, mặc định','Nguồn ứng viên để xếp hạng'),('Đầu ra','Ngày tập, bài tập, sets/reps hoặc thời lượng, nghỉ, tạ','Giáo án cá nhân hóa có thể giải thích')],[.9,2.75,2.85])
doc.add_heading('1.2. Nguyên tắc ưu tiên',2)
doc.add_paragraph('Ràng buộc cứng được áp dụng trước tiêu chí tối ưu: bài không an toàn, thiếu thiết bị hoặc bị người dùng từ chối sẽ không được chọn dù có điểm mục tiêu cao. Sau khi qua bộ lọc, hệ thống mới ưu tiên độ khó và điểm phù hợp mục tiêu.')

doc.add_heading('2. Tính Fitness Score và phân loại thể lực',1)
doc.add_heading('2.1. Công thức Fitness Score',2)
doc.add_paragraph('Điểm thể lực FS nằm trong [0, 100] và là tổng có trọng số của điểm tuổi và điểm cân nặng:')
p=doc.add_paragraph('FS = clamp(w_age × S_age + w_weight × S_weight, 0, 100)'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.runs[0].bold=True
doc.add_paragraph('Mặc định w_age = 0,4 và w_weight = 0,6; hai hệ số được đọc từ SystemConfig nên quản trị viên có thể cấu hình. Nếu thiếu tuổi/chiều cao/cân nặng, hệ thống trả 60; dữ liệu không hữu hạn hoặc không dương trả 0.')
doc.add_heading('2.2. Điểm tuổi',2)
table(doc,['Điều kiện tuổi','Công thức S_age'],[
 ('18-25','100'),('< 18','clamp(100 - (18 - tuổi) × 2, 0, 100)'),('26-40','clamp(100 - (tuổi - 25) × 1,5, 0, 100)'),('> 40','clamp(80 - (min(tuổi, 80) - 40) × 2, 0, 100)')],[2.0,4.5])
doc.add_heading('2.3. Điểm cân nặng',2)
doc.add_paragraph('Cân nặng chuẩn được dùng thống nhất cho FS và BodyType: W_chuẩn = (chiều cao_cm - 100) × 0,9; với nữ nhân thêm 0,9. Độ lệch phần trăm và điểm cân nặng:')
p=doc.add_paragraph('D = |W - W_chuẩn| / W_chuẩn × 100%;     S_weight = clamp(100 - 2D, 0, 100)'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
table(doc,['Khoảng FS','FsLevel','FitnessLevel dùng cho giáo án'],[('85-100','EXCELLENT','ADVANCED'),('65-<85','GOOD','INTERMEDIATE'),('50-<65','AVERAGE','BEGINNER'),('< 50','WEAK','BEGINNER')],[1.4,1.8,3.3])

doc.add_heading('3. Phân loại thể trạng cơ thể (BodyType)',1)
doc.add_paragraph('Thuật toán dùng BMI làm trục chính, kết hợp độ lệch Δ = W - W_chuẩn. Khi BMI ≥ 25 và có phần trăm mỡ cơ thể, % mỡ được ưu tiên để phân biệt khối cơ cao với thừa mỡ.')
table(doc,['Điều kiện','Kết quả'],[
 ('BMI < 18,5 và Δ < -8','CAO_GAY'),('BMI < 18,5 và Δ ≥ -8','GAY_CAN_DOI'),('18,5 ≤ BMI < 25; Δ < -5','GAY_CAN_DOI'),('18,5 ≤ BMI < 25; -5 ≤ Δ < 5','CAN_DOI'),('18,5 ≤ BMI < 25; 5 ≤ Δ ≤ 8','CO_BAP'),('18,5 ≤ BMI < 25; Δ > 8','VAN_DONG_VIEN'),('BMI ≥ 25; không có % mỡ; Δ ≤ 8','CO_BAP'),('BMI ≥ 25; không có % mỡ; Δ > 8','THUA_CAN')],[4.1,2.4])
doc.add_paragraph('Nếu BMI ≥ 25 và có % mỡ: nam <18% hoặc nữ <25% → VẬN ĐỘNG VIÊN; nam 18-24% hoặc nữ 25-31% → CƠ BẮP; cao hơn → THỪA CÂN.',style='Callout')

doc.add_heading('4. Xác định số buổi và mức an toàn',1)
doc.add_paragraph('Số buổi yêu cầu được chặn bởi lịch rảnh và giới hạn theo mục tiêu. Hệ thống không ép người dùng tập nhiều hơn số ngày đã khai báo.')
table(doc,['Mục tiêu','Khoảng cấu hình','Split điển hình'],[
 ('MUSCLE_GAIN','1-6; bảng split chuẩn hỗ trợ 4-6','Push/Pull-Legs lặp lại, bổ sung Full body/Cardio'),('WEIGHT_LOSS','1-6; bảng split chuẩn hỗ trợ 4-6','Ưu tiên Full body/Cardio xen kẽ kháng lực'),('ENDURANCE','1-4; bảng split chuẩn hỗ trợ 2-4','Full body + Cardio mỗi buổi'),('MAINTENANCE','1-5; bảng split chuẩn hỗ trợ 3-5','Cân bằng thân trên, thân dưới, toàn thân')],[1.5,1.65,3.35])
doc.add_paragraph('Điều chỉnh an toàn: với mục tiêu giảm cân, BMI > 30 làm ADVANCED hạ xuống INTERMEDIATE; BMI > 35 chuyển về BEGINNER. Nếu trường trainingExperienceMonths lớn hơn 12 (theo cách dùng hiện tại là thời gian gián đoạn), mức tập giảm một bậc.')

doc.add_heading('5. Phân bổ quota bài tập theo nhóm cơ',1)
doc.add_heading('5.1. BaseQuota',2)
table(doc,['Mục tiêu','BEGINNER','INTERMEDIATE','ADVANCED'],[
 ('Tăng cơ','4','4','6'),('Giảm cân','4','4','6'),('Sức bền','4','4','6'),('Duy trì','3','4','5')],[2.2,1.4,1.4,1.4])
doc.add_heading('5.2. Largest Remainder Method',2)
doc.add_paragraph('Với mỗi nhóm cơ, f là số buổi nhóm đó xuất hiện. Nếu f = 1, quota điều chỉnh bằng min(BaseQuota, 4); nếu f ≥ 2, giữ nguyên BaseQuota. Sau đó:')
p=doc.add_paragraph('base = AdjustedQuota div f;    remainder = AdjustedQuota mod f'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('Mỗi buổi nhận base bài; remainder buổi đầu tiên theo dayIndex nhận thêm 1. Mỗi nhóm cơ tối đa 3 bài trong một ngày. Phần vượt ngưỡng được chuyển sang ngày cùng nhóm cơ đang có ít bài nhất; nếu tất cả đã đầy, phần dư bị bỏ.')
doc.add_heading('5.3. Ví dụ',2)
doc.add_paragraph('Người tập ADVANCED, mục tiêu tăng cơ, nhóm Ngực xuất hiện 2 buổi: quota 6, nên mỗi buổi nhận 3 bài. Nếu một nhóm chỉ xuất hiện 1 buổi: quota bị chặn min(6,4)=4, rồi giới hạn ngày là 3; không có ngày khác để chuyển nên 1 bài dư bị bỏ.')
doc.add_heading('5.4. Giả mã',2)
doc.add_paragraph('''FOR mỗi nhóm cơ g trong tuần:\n    days = các ngày có g\n    f = số phần tử days\n    adjusted = (f == 1) ? min(baseQuota, 4) : baseQuota\n    base = adjusted div f; remainder = adjusted mod f\n    phân phối base + 1 cho remainder ngày đầu\n    chặn mỗi ngày tối đa 3 và chuyển overflow sang ngày ít bài nhất\nRETURN bản đồ ngày -> (nhóm cơ -> số bài)''',style='CodeBlock')

doc.add_heading('6. Lọc và xếp hạng bài tập',1)
doc.add_heading('6.1. Thứ tự độ khó',2)
table(doc,['Trình độ','Thứ tự duyệt ứng viên'],[('BEGINNER','EASY → MEDIUM → HARD'),('INTERMEDIATE','MEDIUM → EASY → HARD'),('ADVANCED','HARD → MEDIUM → EASY')],[2.0,4.5])
doc.add_heading('6.2. Bộ lọc cứng',2)
for x in ['Loại bài có tên khớp danh sách bài người dùng không thích.','Loại bài có chống chỉ định giao với vùng chấn thương; có thêm luật dự phòng theo tên bài cho gối, lưng dưới, vai, cổ tay, khuỷu tay, cổ chân và cổ.','Nếu tập tại nhà, suy luận dụng cụ từ tên bài và chỉ giữ bodyweight hoặc dụng cụ người dùng có. Nếu tập tại GYM/BOTH hoặc danh sách dụng cụ trống, không chặn theo dụng cụ.']:
    doc.add_paragraph(x,style='List Bullet')
doc.add_heading('6.3. Chấm điểm và tính xác định',2)
doc.add_paragraph('Trong từng mức độ khó, ứng viên được sắp giảm dần theo score của mục tiêu. Bằng điểm thì ID nhỏ hơn đứng trước, nhờ đó kết quả ổn định và có thể tái hiện.')
table(doc,['Mục tiêu','Trường điểm'],[('MUSCLE_GAIN','muscleGainScore'),('WEIGHT_LOSS','weightLossScore'),('ENDURANCE','flexibilityScore (tên trường kế thừa)'),('MAINTENANCE','maintenanceScore')],[2.2,4.3])
doc.add_paragraph('Lưu ý kỹ thuật: ENDURANCE hiện sử dụng trường flexibilityScore do di sản mô hình dữ liệu. Tên trường nên được đổi thành enduranceScore trong phiên bản sau để tránh hiểu nhầm.',style='Callout')

doc.add_heading('7. Cá nhân hóa sets, reps, thời lượng và nghỉ',1)
doc.add_heading('7.1. Sets/Reps cơ sở theo FS × mục tiêu',2)
table(doc,['Mục tiêu','EXCELLENT','GOOD','AVERAGE','WEAK'],[
 ('Tăng cơ','4×6','4×8','3×10','3×12'),('Giảm cân','4×15','4×14','3×13','3×12'),('Sức bền','3×15','3×13','2×12','2×10'),('Duy trì','3×10','3×11','3×12','3×13')],[1.5,1.25,1.25,1.25,1.25])
doc.add_heading('7.2. Điều chỉnh theo BodyType',2)
doc.add_paragraph('Hệ thống cộng vector điều chỉnh [Δsets, Δreps] theo BodyType và mục tiêu. Ví dụ: tăng cơ + CAO_GAY → [0,+2]; tăng cơ + THUA_CAN → [+1,0]; giảm cân + THUA_CAN → [+2,0]; duy trì + VAN_DONG_VIEN → [0,-1].')
doc.add_heading('7.3. Training Zone và gợi ý tải',2)
table(doc,['Mục tiêu','Rep floor','Rep ceiling'],[('Tăng cơ','5','12'),('Giảm cân','12','20'),('Sức bền','8','15'),('Duy trì','8','15')],[2.5,2.0,2.0])
doc.add_paragraph('finalSets = max(1, baseSets + Δsets). rawRep = baseReps + Δreps và finalReps = clamp(rawRep, floor, ceiling). Nếu rawRep vượt trần, hệ thống giữ rep ở trần và gợi ý tăng tạ; nếu thấp hơn sàn thì gợi ý giảm tạ.')
doc.add_heading('7.4. Thời lượng và nghỉ',2)
table(doc,['Thành phần','Quy tắc mặc định'],[('Thời lượng bài','BEGINNER ×0,7; INTERMEDIATE ×1,0; ADVANCED ×1,3'),('Nghỉ - tăng cơ','Thời gian nghỉ gốc ×1,3'),('Nghỉ - giảm cân','Thời gian nghỉ gốc ×0,7'),('Nghỉ - sức bền/duy trì','Giữ thời gian nghỉ gốc'),('Giới hạn bài/buổi','maxExercises = clamp(durationMinutes div 10, 2, 10)')],[2.3,4.2])

doc.add_heading('8. Tính mức tạ khuyến nghị',1)
doc.add_paragraph('Chỉ áp dụng cho bài usesWeight=true và các nhóm Ngực, Lưng, Chân, Vai, Tay, Core. Công thức:')
p=doc.add_paragraph('Load = round0.5(W × F_muscle × F_FS × F_body × F_goal × F_delta)'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.runs[0].bold=True
table(doc,['Hệ số','Giá trị'],[
 ('F_muscle','Ngực 0,40; Lưng 0,50; Chân 0,60; Vai 0,25; Tay 0,20; Core 0,15'),('F_FS','≥90: 1,20; ≥80: 1,10; ≥60: 1,00; ≥40: 0,85; còn lại: 0,75'),('F_body','Cao gầy 0,85; gầy cân đối 0,90; cân đối 1,00; cơ bắp 1,10; VĐV 1,20; thừa cân 0,95'),('F_goal','Tăng cơ 1,05; giảm cân 0,90; duy trì 1,00; sức bền 0,60'),('F_delta','Δ<-10:0,85; <-5:0,90; ≤5:1,00; ≤10:1,05; >10:1,00')],[1.25,5.25])
doc.add_paragraph('Kết quả được làm tròn đến 0,5 kg. Đây là mức khởi điểm tham khảo; người tập vẫn cần ưu tiên kỹ thuật và phản hồi thực tế.',style='Callout')

doc.add_heading('9. Thuật toán tổng thể',1)
doc.add_paragraph('''INPUT: profile, goal, requestedLevel, requestedSessions, exerciseCatalog\n1. fs = calculateFS(age, height, weight, gender)\n2. fsLevel = classifyFS(fs); bodyType = classifyBodyType(profile)\n3. safeLevel = adjustByBMIAndExperience(requestedLevel)\n4. sessions = clampByGoalAndAvailability(requestedSessions)\n5. weekPlan = splitAndDistribute(goal, safeLevel, sessions)\n6. FOR mỗi ngày và mỗi nhóm cơ:\n       pool = query(active, muscleGroup, difficulty)\n       pool = filter(disliked, injuries, equipment)\n       pool = sort(goalScore DESC, id ASC)\n       chọn đủ quota\n       gán sets/reps, duration, rest, recommendedLoad\n7. cắt danh sách theo thời lượng buổi\nOUTPUT: workoutPlan''',style='CodeBlock')
doc.add_heading('9.1. Độ phức tạp',2)
doc.add_paragraph('Gọi N là tổng số bài ứng viên, D là số buổi và G là số nhóm cơ. Phân bổ quota có độ phức tạp xấp xỉ O(D×G). Lọc là O(N), sắp xếp theo từng pool là O(N log N) trong trường hợp xấu nhất. Bộ nhớ phụ chủ yếu O(N) cho danh sách ứng viên.')

doc.add_heading('10. Ví dụ minh họa',1)
doc.add_heading('10.1. Hồ sơ tăng cơ',2)
doc.add_paragraph('Nam, 25 tuổi, cao 175 cm, nặng 70 kg, mục tiêu tăng cơ, trình độ trung cấp, 4 buổi/tuần, tập tại gym. W_chuẩn=67,5 kg; D≈3,70%; S_weight≈92,59; S_age=100; FS≈95,56 → EXCELLENT/ADVANCED. BodyType=CÂN ĐỐI. Sets/reps cơ sở 4×6, không điều chỉnh BodyType, nằm trong zone 5-12. Split 4 buổi phân các nhóm Ngực/Vai/Tay và Lưng/Core/Chân lặp hai lần; quota ADVANCED=6 nên mỗi nhóm nhận 3 bài ở mỗi lần xuất hiện.')
doc.add_heading('10.2. Hồ sơ giảm cân có chấn thương',2)
doc.add_paragraph('Người mới tập tại nhà, BMI > 35, 45 phút/buổi, đau gối và không thích Burpee. Mức tập được hạ về BEGINNER. Bộ lọc loại Squat, Lunge, Leg Press, bài nhảy và Burpee; đồng thời loại bài cần thiết bị không có. Các bài còn lại được xếp theo weightLossScore. Mỗi buổi tối đa clamp(45 div 10,2,10)=4 bài. Sets/reps được tính theo FS/BodyType nhưng giữ trong zone 12-20 reps.')

doc.add_heading('11. Tính đúng đắn, giới hạn và hướng cải tiến',1)
doc.add_heading('11.1. Thuộc tính bảo đảm',2)
for x in ['Tính xác định: tie-break theo ID, không dùng random.','An toàn ưu tiên: lọc cứng trước chấm điểm.','Giới hạn tải: training zone, tối thiểu 1 set, tối đa 3 bài/nhóm cơ/ngày.','Khả năng giải thích: mọi kết quả truy ngược được về bảng luật và hệ số cấu hình.','Khả năng cấu hình: trọng số FS, multiplier thời lượng/nghỉ và lịch khuyến nghị được tách khỏi logic giao diện.']:
    doc.add_paragraph(x,style='List Bullet')
doc.add_heading('11.2. Giới hạn hiện tại',2)
for x in ['Suy luận dụng cụ và luật chấn thương dự phòng dựa trên tên bài, phụ thuộc cách đặt tên.','Nếu danh sách thiết bị trống khi tập tại nhà, logic hiện tại cho phép mọi bài; nên xác định rõ “trống” là chưa khai báo hay không có thiết bị.','Cắt danh sách theo thời lượng xảy ra sau khi ghép nhóm cơ, có thể làm các nhóm cuối bị thiếu bài.','Công thức mức tạ là heuristic theo cân nặng cơ thể, chưa dùng 1RM/RPE hoặc lịch sử hoàn thành.','Trường trainingExperienceMonths đang được dùng như thời gian gián đoạn >12 tháng; cần đổi tên hoặc làm rõ nghĩa dữ liệu.','Các ngưỡng BMI/FS là luật nghiệp vụ, không thay thế đánh giá chuyên môn y tế.']:
    doc.add_paragraph(x,style='List Bullet')
doc.add_heading('11.3. Hướng cải tiến',2)
for x in ['Chuẩn hóa requiredEquipment và contraindicatedInjuries thành dữ liệu cấu trúc bắt buộc.','Cân bằng lại bước cắt theo thời lượng bằng round-robin giữa các nhóm cơ.','Bổ sung progressive overload từ RPE, tỷ lệ hoàn thành và mức tạ thực tế.','Đổi flexibilityScore thành enduranceScore và migration dữ liệu.','Mở rộng kiểm thử biên cho BMI, FS, quota, chấn thương và thiết bị.']:
    doc.add_paragraph(x,style='List Number')

doc.add_heading('12. Đối chiếu mã nguồn',1)
table(doc,['Thành phần','Tệp mã nguồn','Trách nhiệm'],[
 ('FitnessCalculator','service/FitnessCalculator.java','FS, FsLevel, BodyType, sets/reps và clamp zone'),('MuscleGroupSplitPlanner','service/plan/MuscleGroupSplitPlanner.java','Split tuần, quota, LRM, cap và redistribution'),('TrainingZone','service/setrep/TrainingZone.java','Khoảng reps theo mục tiêu'),('WorkoutPlanService','service/WorkoutPlanService.java','Điều phối, lọc, score, duration, rest và mức tạ'),('TrainingConfigService','service/TrainingConfigService.java','Lịch và cấu hình nhóm cơ có thể quản trị')],[1.55,2.65,2.3])
doc.add_paragraph('Kết luận: thuật toán GymPro là hệ luật nhiều tầng, trong đó an toàn và khả năng thực hiện là điều kiện bắt buộc; mục tiêu và thể lực quyết định mức ưu tiên và khối lượng; thời lượng buổi là giới hạn cuối cùng. Cấu trúc này phù hợp với đồ án vì dễ kiểm thử, giải thích và mở rộng.',style='Callout')

OUT.parent.mkdir(parents=True,exist_ok=True)
doc.core_properties.title='Tài liệu thuật toán sinh giáo án cá nhân hóa GymPro'; doc.core_properties.subject='Đồ án tốt nghiệp GymPro'; doc.core_properties.author='Nhóm phát triển GymPro'
doc.save(OUT)
print(OUT)
